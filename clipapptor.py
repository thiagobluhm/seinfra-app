import gradio as gr
import os
os.chdir(os.path.abspath(os.curdir))
import time
import fitz  # PyMuPDF
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tqdm import tqdm
from dotenv import load_dotenv
from modelo_clip import ValidatorVITImgTexto
from openai import OpenAI
import sys
import io
import re
from threading import Thread
from queue import Queue


# Capturador de logs
class StreamCapturer(io.StringIO):
    def __init__(self, queue):
        super().__init__()
        self.queue = queue

    def write(self, s):
        super().write(s)
        self.queue.put(s)

# Criar fila para logs e capturar saída
log_queue = Queue()
sys.stdout = StreamCapturer(log_queue)

# Função para atualizar logs em tempo real
def monitor_logs():
    """ Atualiza a interface do Gradio com os logs capturados """
    log_text = ""
    while not log_queue.empty():
        try:
            log_text += log_queue.get(timeout=1) + "\n"
            #log_output.update(log_text)  # Atualiza log na interface
            return log_text
        except:
            pass

# Carregar variáveis de ambiente
load_dotenv()
secretk = os.environ.get("OPENAI_API_KEY")
cliente_openai = OpenAI(api_key=secretk)

# Criar pastas para armazenar arquivos
os.makedirs("uploads", exist_ok=True)
os.makedirs("clippings", exist_ok=True)

# 🔹 Justificar parágrafos no Word
def justify_paragraph(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)

# 🔹 Função para extrair texto e imagens do PDF
def extract_content(pdf_path):
    doc = fitz.open(pdf_path)
    extracted_data = []

    for page_num, page in tqdm(enumerate(doc, start=1), desc="Extraindo páginas do PDF", unit="página"):
        page_data = {"page_number": page_num, "text": page.get_text("text"), "images": []}
        
        for _, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_data = base_image["image"]
            try:
                image = Image.open(BytesIO(image_data))
                page_data["images"].append(image)
            except Exception as e:
                print(f"Erro ao processar imagem na página {page_num}: {e}")

        extracted_data.append(page_data)
    return extracted_data


def sanitize_text(text):
    """
    Remove caracteres não suportados pelo formato XML usado no Word.
    """
    # Remove caracteres de controle e bytes nulos
    sanitized_text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
    
    # Garante que o texto seja Unicode
    return sanitized_text.encode('utf-8', 'ignore').decode('utf-8')

# 🔹 Verifica se um texto contém palavras-chave relevantes
def is_relevant(text, keywords):
    return any(keyword.lower() in text.lower() for keyword in keywords)

# 🔹 Filtrar conteúdo com base no prompt e nas palavras-chave
def filter_content(prompt, extracted_data, keywords):
    filtered_data = []

    for page in tqdm(extracted_data, desc="Classificando páginas com OpenAI", unit="página"):
        if page["text"].strip():
            try:
                response = cliente_openai.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "Você é um assistente que classifica textos extraídos de PDFs."},
                        {"role": "user", "content": f"Prompt: {prompt}\nTexto: {page['text']}"}
                    ], 
                    temperature=0.3, 
                    max_tokens=4096
                )
                relevance = response.choices[0].message.content
                if "relevante" in relevance.lower() or is_relevant(page["text"], keywords):
                    filtered_data.append(page)

            except Exception as e:
                print(f"Erro ao classificar conteúdo na página {page['page_number']}: {e}")
        elif page["images"]:  # Se não há texto, mas há imagens, ainda incluímos
            filtered_data.append(page)

    return filtered_data

# 🔹 Geração do Word com os textos e imagens filtrados
def generate_word_from_prompt(filtered_data, output_path, validator):
    doc = Document()

    if not filtered_data:
        doc.add_paragraph("Nenhum conteúdo relevante encontrado.")
    else:
        for page in tqdm(filtered_data, desc="Adicionando páginas ao Word", unit="página"):
            doc.add_heading(f"Notícia (Página {page['page_number']})", level=2)

            for image in page["images"]:
                truncated_texts = [validator.truncate_text(page["text"])]
                top_texts = validator.get_top_k_texts(image, truncated_texts, k=3)

                if top_texts:
                    img_temp_path = f"temp_image_{time.time()}.png"
                    image.save(img_temp_path)
                    doc.add_picture(img_temp_path, width=Inches(5))
                    os.remove(img_temp_path)

                    for text, score in top_texts:
                        text = sanitize_text(text)
                        paragraph = doc.add_paragraph(f"Relevância: {score:.2f}\n{text}")
                        justify_paragraph(paragraph)

            if page["text"].strip():
                text = sanitize_text(page["text"])
                paragraph = doc.add_paragraph(text)
                justify_paragraph(paragraph)

            doc.add_paragraph("-" * 50)

    doc.save(output_path)

# 🔹 Processamento do PDF enviado pelo Gradio
def process_pdf(uploaded_pdf):
    if not uploaded_pdf:
        return "Nenhum arquivo enviado.", None

    # 🔥 **Correção do erro: Gradio envia caminho do arquivo, então apenas pegamos o `name`**
    pdf_path = uploaded_pdf#.name  # Caminho do arquivo carregado
    output_docx = os.path.join("clippings", f"clipping_{int(time.time())}.docx")

    # Inicializar Validator
    validator = ValidatorVITImgTexto()
    print("📌 Iniciando processamento do PDF...")
    #print(f"📂 Arquivo recebido: {pdf_path}")
    # Prompt de exemplo
    prompt = """
                Você é o analista responsável pelo Clipping do Órgão SEINFRA-CE. 
                Sua tarefa é Extrair todas as notícias sobre:

                1. Governo do Estado do Ceará: qualquer assunto, incluindo saúde, infraestrutura, educação, segurança pública e política.
                2. Governo Federal: notícias de infraestrutura, saúde, política nacional e temas relevantes para órgãos públicos.
                3. Política Nacional: temas relacionados a governadores, prefeitos, ministros, deputados e senadores.
                4. Infraestrutura e Saúde: qualquer projeto, evento ou decisão relacionada a esses temas.

                Não deixe de pegar nenhum texto importante. Priorize textos que mencionem 'Governo Federal', 'Governo do Estado do Ceará' ou seus representantes.
            """
    keywords = [
        "Governo Federal", "Governo do Estado do Ceará", "Ceará", "infraestrutura",
        "saúde", "educação", "segurança pública", "investimento", "ministro", 
        "governador", "prefeito", "política", "infraestrutura urbana", "mobilidade urbana",
        "transporte público", "energia renovável", "obras públicas", "hospital", 
        "escolas públicas", "recursos federais", "recursos estaduais", "moradia", 
        "habitação", "economia", "desenvolvimento regional", "assistência social", 
        "meio ambiente", "água potável", "saneamento básico", "pontes", "rodovias", 
        "ferrovias", "aeroportos", "portos", "logística", "segurança alimentar", 
        "orçamento público", "políticas públicas", "emendas parlamentares", 
        "licitação", "contratos públicos", "parcerias público-privadas", 
        "inovação tecnológica", "tecnologia", "saúde pública", "sistema único de saúde",
        "vacinação", "doenças infecciosas", "crise hídrica", "ministério", 
        "agricultura", "turismo", "indústria", "comércio exterior", "universidades"
    ]

    print("🔍 Extraindo conteúdo do PDF...")
    extracted_data = extract_content(pdf_path)

    print("📑 Classificando textos extraídos...")
    filtered_data = filter_content(prompt, extracted_data, keywords)

    print("✍️ Gerando documento final...")
    generate_word_from_prompt(filtered_data, output_docx, validator)

    print("✅ Processo concluído! Arquivo salvo.")
    
    return gr.update(visible=True), output_docx, monitor_logs() #"✅ Processamento concluído! Verifique o resultado."

from gradio_pdf import PDF
# Interface do Gradio
with gr.Blocks() as demo:
    gr.Markdown("### 📑 Clipapptor - IA*")
    
    with gr.Row():
        with gr.Column(scale=1):  # Lado esquerdo

            uploaded_pdf = PDF(label="📤 Faça upload do PDF", interactive=True, height=650)
                   
            
        with gr.Column(scale=1):  # Lado direito
            name = gr.Textbox(label = "📄 Caminho do arquivo.")
            uploaded_pdf.upload(lambda f: f, inputs=uploaded_pdf, outputs=name)
            #output_pdf_view = gr.HTML()  # Visualização do PDF         
            process_button = gr.Button("📄 Processar PDF", interactive=True)     
            #log_output = gr.Textbox(label="🖥️ Terminal", lines=15, interactive=False)  # Terminal            
            output_docx_view = gr.File(label="📋 Baixar Clipping Gerado")    

        #uploaded_pdf.change(fn=update_pdf_view, inputs=[uploaded_pdf], outputs=[output_pdf_view])
        process_button.click(fn=process_pdf, inputs=[uploaded_pdf], outputs=[uploaded_pdf, output_docx_view])

# Iniciar monitor de logs em background
# log_thread = Thread(target=monitor_logs, args=(log_output,), daemon=True)
# log_thread.start()

# Executar interface
demo.launch()



    