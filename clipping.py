from openai import OpenAI
from modelo_clip import ValidatorVITImgTexto
import fitz  # PyMuPDF
from PIL import Image
from io import BytesIO
import pytesseract
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tqdm import tqdm  # Importação do tqdm
import time  # Para contagem de tempo
from dotenv import load_dotenv
load_dotenv()
import os
os.chdir(os.path.abspath(os.curdir))

secretk = os.environ.get("OPENAI_API_KEY")
cliente_openai = OpenAI(api_key=secretk)

# Função para justificar o texto e evitar quebras de palavras
def justify_paragraph(paragraph):
    # Garantir que o parágrafo seja justificado
    p = paragraph._element  # Acessar o elemento XML do parágrafo
    pPr = p.get_or_add_pPr()  # Obter ou criar as propriedades do parágrafo
    jc = OxmlElement('w:jc')  # Criar o elemento de justificação
    jc.set(qn('w:val'), 'both')  # Definir a justificação como 'both' (justificado)
    pPr.append(jc)
    
    # Adicionar configuração para evitar hifenização
    no_hyphenation = OxmlElement('w:noHyphenation')
    pPr.append(no_hyphenation)

# Adicionar validação de imagem e texto usando ValidatorVITImgTexto
def validate_image_with_text(validator, image, text):
    """
    Valida a relevância da imagem em relação ao texto usando ValidatorVITImgTexto.
    """
    relevance_score = validator.validate(image, text)
    return relevance_score > 0.5  # Retorna True se a pontuação for alta

# Função para extrair texto e imagens do PDF
def extract_content(pdf_path):
    doc = fitz.open(pdf_path)
    extracted_data = []  # Lista para armazenar texto e imagens por página

    for page_num, page in tqdm(enumerate(doc, start=1), desc="Extraindo páginas do PDF", unit="página"):
        page_data = {"page_number": page_num, "text": page.get_text("text"), "images": []}

        # Extrair imagens
        for _, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_data = base_image["image"]  # Dados binários da imagem
            try:
                image = Image.open(BytesIO(image_data))  # Converter binário para imagem
                page_data["images"].append(image)
            except Exception as e:
                print(f"Erro ao processar imagem na página {page_num}: {e}")

        extracted_data.append(page_data)
    return extracted_data

# Filtro de palavras-chave
def is_relevant(text, keywords):
    for keyword in keywords:
        if keyword.lower() in text.lower():
            return True
    return False

def filter_content(prompt, extracted_data, keywords):
    filtered_data = []

    for page in tqdm(extracted_data, desc="Classificando páginas com OpenAI", unit="página"):
        if page["text"].strip():  # Verifica se há texto na página

            try:
                response = cliente_openai.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "Você é um assistente que classifica textos extraídos de PDFs."},
                        {"role": "user", "content": f"Prompt: {prompt}\nTexto: {page['text']}"}
                    ], 
                    temperature=0.3, 
                    max_tokens=4096
                )
                relevance = response.choices[0].message.content
                if "relevante" in relevance.lower() or is_relevant(page["text"], keywords):
                    filtered_data.append(page)

            except Exception as e:
                print(f"Erro ao classificar conteúdo na página {page['page_number']}: {e}")
        
        else:
            # Adicionar página sem texto relevante apenas com imagens
            if page["images"]:
                filtered_data.append(page)

    return filtered_data

def generate_word_from_prompt(filtered_data, output_path, validator):
    doc = Document()

    if not filtered_data:
        doc.add_paragraph("Nenhum conteúdo relevante encontrado.")
    else:
        for page in tqdm(filtered_data, desc="Adicionando páginas ao Word", unit="página"):
            doc.add_heading(f"Notícia (Página {page['page_number']})", level=2)

            # Processar imagens e validar relação com o texto
            for image in page["images"]:
                truncated_texts = [validator.truncate_text(page["text"])]  # Garante truncagem
                top_texts = validator.get_top_k_texts(image, truncated_texts, k=3)

                if top_texts:
                    img_temp_path = f"temp_image_{time.time()}.png"
                    image.save(img_temp_path)
                    doc.add_picture(img_temp_path, width=Inches(5))
                    os.remove(img_temp_path)

                    # Adicionar textos relevantes abaixo da imagem
                    for text, score in top_texts:
                        paragraph = doc.add_paragraph(f"Relevância: {score:.2f}\n{text}")
                        justify_paragraph(paragraph)

            # Adicionar texto principal, se houver
            if page["text"].strip():
                paragraph = doc.add_paragraph(page["text"])
                justify_paragraph(paragraph)

            # Separador
            doc.add_paragraph("-" * 50)

    doc.save(output_path)


# # Função principal
# def main():
#     pdf_path = r"pdfs/OPI 21.11.pdf"  # Substituir pelo caminho do seu PDF
#     output_path = r"clippings/resultado.docx"
    
#     # Inicializar o Validator
#     validator = ValidatorVITImgTexto()

    # # Prompt de exemplo
    # prompt = """
    #             Você é o analista responsável pelo Clipping do Órgão SEINFRA-CE. 
    #             Sua tarefa é Extrair todas as notícias sobre:

    #             1. Governo do Estado do Ceará: qualquer assunto, incluindo saúde, infraestrutura, educação, segurança pública e política.
    #             2. Governo Federal: notícias de infraestrutura, saúde, política nacional e temas relevantes para órgãos públicos.
    #             3. Política Nacional: temas relacionados a governadores, prefeitos, ministros, deputados e senadores.
    #             4. Infraestrutura e Saúde: qualquer projeto, evento ou decisão relacionada a esses temas.

    #             Não deixe de pegar nenhum texto importante. Priorize textos que mencionem 'Governo Federal', 'Governo do Estado do Ceará' ou seus representantes.
    #         """
    # keywords = [
    #     "Governo Federal", "Governo do Estado do Ceará", "Ceará", "infraestrutura",
    #     "saúde", "educação", "segurança pública", "investimento", "ministro", 
    #     "governador", "prefeito", "política", "infraestrutura urbana", "mobilidade urbana",
    #     "transporte público", "energia renovável", "obras públicas", "hospital", 
    #     "escolas públicas", "recursos federais", "recursos estaduais", "moradia", 
    #     "habitação", "economia", "desenvolvimento regional", "assistência social", 
    #     "meio ambiente", "água potável", "saneamento básico", "pontes", "rodovias", 
    #     "ferrovias", "aeroportos", "portos", "logística", "segurança alimentar", 
    #     "orçamento público", "políticas públicas", "emendas parlamentares", 
    #     "licitação", "contratos públicos", "parcerias público-privadas", 
    #     "inovação tecnológica", "tecnologia", "saúde pública", "sistema único de saúde",
    #     "vacinação", "doenças infecciosas", "crise hídrica", "ministério", 
    #     "agricultura", "turismo", "indústria", "comércio exterior", "universidades"
    # ]

#     # Etapas do pipeline com contagem de tempo
#     start_time = time.time()

#     print("Extraindo conteúdo do PDF...")
#     extract_start = time.time()
#     extracted_data = extract_content(pdf_path)
#     extract_end = time.time()
#     print(f"Conteúdo extraído em {extract_end - extract_start:.2f} segundos.")

#     print("Classificando textos extraídos com base no prompt...")
#     filter_start = time.time()
#     filtered_data = filter_content(prompt, extracted_data, keywords)
#     filter_end = time.time()
#     print(f"Classificação concluída em {filter_end - filter_start:.2f} segundos.")

#     print("Gerando o documento Word com os textos e imagens filtrados...")
#     generate_start = time.time()
#     generate_word_from_prompt(filtered_data, output_path, validator)
#     generate_end = time.time()
#     print(f"Documento gerado em {generate_end - generate_start:.2f} segundos.")

#     end_time = time.time()
#     print(f"Processo concluído! Documento salvo em: {output_path}")
#     print(f"Tempo total de execução: {end_time - start_time:.2f} segundos.")

# # Chamar a função principal
# if __name__ == "__main__":
#     main()